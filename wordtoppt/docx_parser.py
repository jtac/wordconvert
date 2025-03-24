"""
DocxParser: Extract and structure content from Word documents.
"""

import logging
from typing import Dict, List, Any

import docx

logger = logging.getLogger(__name__)


class DocxParser:
    """
    Extracts and structures content from Word documents.
    """

    def __init__(self, file_path: str):
        """
        Initialize the DocxParser.

        Args:
            file_path: Path to the Word document
        """
        self.file_path = file_path
        self.document = None

    def extract_content(self) -> Dict[str, Any]:
        """
        Extract content from the Word document.

        Returns:
            Dict containing structured document content
        """
        try:
            self.document = docx.Document(self.file_path)
        except Exception as e:
            logger.error("Error opening document: %s", str(e))
            raise RuntimeError(f"Failed to open document: {str(e)}") from e

        # Extract document structure
        content = {
            "title": self._extract_title(),
            "headings": self._extract_headings(),
            "paragraphs": self._extract_paragraphs(),
            "sections": self._extract_sections(),
        }

        return content

    def _extract_title(self) -> str:
        """Extract the document title."""
        if not self.document or not self.document.paragraphs:
            return ""

        # Attempt to find title - typically the first paragraph with large font or Heading 1
        for paragraph in self.document.paragraphs[:3]:  # Check first few paragraphs
            if (
                paragraph.style.name.startswith("Heading 1")
                or paragraph.style.name == "Title"
            ):
                return paragraph.text.strip()

        # Fallback to first non-empty paragraph
        for paragraph in self.document.paragraphs[:5]:
            if paragraph.text.strip():
                return paragraph.text.strip()

        return ""

    def _extract_headings(self) -> List[Dict[str, Any]]:
        """Extract all headings and their hierarchy."""
        headings = []

        if not self.document:
            return headings

        for paragraph in self.document.paragraphs:
            if paragraph.style.name.startswith("Heading"):
                level = int(paragraph.style.name.split()[-1])
                headings.append(
                    {
                        "level": level,
                        "text": paragraph.text.strip(),
                    }
                )

        return headings

    def _extract_paragraphs(self) -> List[str]:
        """Extract all paragraphs."""
        paragraphs = []

        if not self.document:
            return paragraphs

        for paragraph in self.document.paragraphs:
            # Skip headings and empty paragraphs
            if (
                not paragraph.style.name.startswith("Heading")
                and paragraph.text.strip()
            ):
                paragraphs.append(paragraph.text.strip())

        return paragraphs

    def _extract_sections(self) -> List[Dict[str, Any]]:
        """
        Extract document sections based on headings.
        Each section contains its heading and all content until the next heading.
        """
        sections = []
        current_section = None
        current_content = []

        if not self.document:
            return sections

        for paragraph in self.document.paragraphs:
            if paragraph.style.name.startswith("Heading"):
                # Save previous section if exists
                if current_section:
                    sections.append(
                        {
                            "heading": current_section,
                            "content": current_content,
                        }
                    )

                # Start new section
                level = int(paragraph.style.name.split()[-1])
                current_section = {
                    "level": level,
                    "text": paragraph.text.strip(),
                }
                current_content = []
            elif paragraph.text.strip() and current_section:
                current_content.append(paragraph.text.strip())

        # Add the last section
        if current_section:
            sections.append(
                {
                    "heading": current_section,
                    "content": current_content,
                }
            )

        return sections
